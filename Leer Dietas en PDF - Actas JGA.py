# -*- coding: utf-8 -*-
"""
=============================================================================
 EXTRACCION DE REMUNERACION DE DIRECTORIOS Y COMITES  ·  v4
 (mas rapido y mas preciso; soporta Word .docx)
=============================================================================

NOVEDADES v4 (tras analizar 269 PDFs reales, 5 horas de corrida):

 ALCANCE / FILTROS
  - Filtro de AÑOS editable (ANIO_MIN..ANIO_MAX): solo procesa las carpetas de
    año pedidas (por defecto 2020..2026).
  - Filtro de paginas DESACTIVADO por defecto (MIN_PAGINAS_ACTA=0): se procesan
    todos los archivos, incluso los de 2 paginas (la info puede venir corta).
    Editable: si subes el numero, vuelve a filtrar actas largas.

 NOMBRES DE EMPRESA = PORTAFOLIO
  - El output usa los nombres EXACTOS de tu portafolio. Mapa de alias/rebrands
    (Edelnor->Pluz, Edegel->Orygen, Enersur->Engie, GyM->Aenza, Milpo/Atacocha->
    Nexa, Continental->BBVA, BAP->Credicorp, HOC->Hochshild, HBM->Hudbay, ...).
    Lo que no es de tu portafolio sale '[SIN MATCH PORTAFOLIO]' (o se omite si
    SOLO_PORTAFOLIO=True). Nunca se inventa una equivalencia.

 VELOCIDAD
  - Paralelismo: varios archivos a la vez (ThreadPoolExecutor). El subproceso de
    Tesseract corre en paralelo de verdad.

 PRECISION
  - Tope de monto: una dieta de director NO es de millones. Se descartan
    cifras > MONTO_MAX (utilidades, dividendos, capital) y < MONTO_MIN.
  - Filtro anti-financiero: utilidades/dividendos/capital/reserva -> no es dieta.

 COBERTURA
  - Soporta Word .docx (parrafos + tablas, sin OCR). Los .doc binarios antiguos
    se marcan para conversion manual.

=============================================================================
 EXTRACCION DE REMUNERACION (base v3: render PyMuPDF, sin poppler)
=============================================================================

DEPENDENCIAS (Windows):
    pip install pymupdf pytesseract opencv-python pillow rapidfuzz pandas openpyxl numpy
    - Tesseract OCR + idioma español (spa.traineddata).  <-- UNICA dependencia externa
        https://github.com/UB-Mannheim/tesseract/wiki
      Ajusta TESSERACT_CMD abajo a tu instalacion.
    - YA NO se necesita poppler (PyMuPDF renderiza el PDF internamente).

Cambio v3:
    - Motor de render: pdf2image/poppler  ->  PyMuPDF (fitz). Sin binarios
      externos ni edicion del PATH.
    - Deteccion de capa de texto NATIVA: si una pagina trae texto real
      (no escaneada), se usa directo y se evita el OCR.

Cambios heredados de v2 (motivados por actas reales Engie/Volcan):

 [1] PAGINAS DE EEFF: las actas incluyen Estados Financieros (Situacion,
     Resultados, Flujos de Efectivo). Lineas como "Pagos a proveedores..."
     generaban falsos positivos. -> Se detectan y se OMITEN esas paginas.

 [2] KEYWORDS FUERTES vs DEBILES: "pago"/"compensacion" solo cuentan si hay
     contexto de directorio/comite/sesion/dieta cerca. "dieta"/"retribucion"/
     "remuneracion"/"honorario" disparan por si solas.

 [3] EXTRACCION POR VENTANA (no por split de "."): se busca la keyword y se
     toma una ventana de caracteres alrededor. Evita romper decimales
     (S/ 1,200.50 ya no se parte en 1200 + 50).

 [4] OCR conserva estructura de lineas (image_to_data por bloque/linea).

 [5] FRECUENCIA tolerante: "por cada sesion", "por sesion presencial", etc.

 [6] GUARDA DE CENTAVOS: ignora la notacion "00/100" deletreada.

 [7] PRIORIDAD DE MONEDA: si en la ventana hay montos con S//US$, se
     descartan los numeros sin moneda (matan el "100" de "00/100").

 [8] EMPRESA: regex en mayuscula+minuscula terminada en S.A./S.A.A./S.A.C.,
     desambiguada con el nombre del archivo. Capta "Engie Energia Peru S.A."
     y "Volcan Compania Minera S.A.A.".

 [9] AÑO: carpeta -> nombre de archivo -> fecha del propio acta.

 [10] HOJA 'DIAGNOSTICO': lista TODOS los archivos procesados y si arrojaron
      datos (asi sabes que Volcan se proceso y dio 0, sin inventar nada).

Restricciones mantenidas: sin xlwings/win32com, openpyxl, comentarios en
español, estructura por celdas '# %%', evidencia textual en cada fila.
=============================================================================
"""

# %% ============================ IMPORTS ====================================
import os
import re
import json
import hashlib
import logging
import unicodedata
from pathlib import Path
from datetime import datetime
from collections import Counter
from concurrent.futures import ThreadPoolExecutor, as_completed

import numpy as np
import pandas as pd

import fitz  # PyMuPDF: renderiza el PDF sin poppler
import pytesseract
from PIL import Image

try:
    import docx  # python-docx, para leer .docx
    _HAY_DOCX = True
except Exception:
    _HAY_DOCX = False

try:
    import cv2
    _HAY_CV2 = True
except Exception:
    _HAY_CV2 = False

from rapidfuzz import fuzz

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


# %% ========================== CONFIGURACION ================================

RUTA_RAIZ = r"Z:\Mesa de Inversiones\Bottom-Up\0 Juntas de Accionistas\Actas"
RUTA_SALIDA_EXCEL = os.path.join(
    RUTA_RAIZ, f"Remuneraciones_Directorio_{datetime.now():%Y%m%d_%H%M}.xlsx"
)
RUTA_CACHE_OCR = os.path.join(RUTA_RAIZ, "_ocr_cache")

TESSERACT_CMD = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# Carpeta que CONTIENE spa.traineddata. Dejar "" si ya esta en la carpeta
# tessdata por defecto de Tesseract. Util si no tienes permisos de admin para
# escribir en "C:\Program Files\Tesseract-OCR\tessdata".
TESSDATA_DIR = r""   # ej: r"C:\Users\usuario\OneDrive\Desktop\AFP INTEGRA\tessdata"

OCR_IDIOMA = "spa"
OCR_DPI = 300
OCR_CONFIG = "--oem 1 --psm 6"
UMBRAL_UPSCALE_PX = 1700
# Si una pagina trae al menos estos caracteres de texto nativo, se usa ese
# texto y se evita el OCR (mas rapido y exacto para PDFs no escaneados).
UMBRAL_TEXTO_NATIVO = 200

CONFIANZA_FILA_REVISION = 70.0    # filas con confianza OCR < esto -> revisar
UMBRAL_SIMILITUD_EMPRESA = 88     # token_set_ratio para fusionar empresas
EMPRESA_DESDE = "filename"        # "filename" (recomendado) | "auto" | "contenido"

# --- FILTRO DE AÑOS (carpetas a procesar). Editable: cambia el rango. ---
ANIO_MIN = 2020
ANIO_MAX = 2026
ANIOS_INCLUIR = set(range(ANIO_MIN, ANIO_MAX + 1))   # {2020, 2021, ..., 2026}

# --- VELOCIDAD ---
# Filtro de paginas DESACTIVADO (=0): se procesan todos los archivos, incluso
# los de 2 paginas (la info puede venir en archivos cortos). Si algun dia
# quieres volver a filtrar actas largas, sube este numero (ej. 8).
MIN_PAGINAS_ACTA = 0
# Numero de archivos a procesar en paralelo. Tesseract corre en subproceso, asi
# que los hilos dan paralelismo real. Baja este numero si la PC se satura.
MAX_WORKERS = max(2, (os.cpu_count() or 4) - 1)
PROCESAR_WORD = True              # incluir archivos .docx

# --- PRECISION: topes de monto (una dieta de director no es de millones) ---
MONTO_MIN = 100.0                 # por debajo: numero de pagina/articulo -> descartar
MONTO_MAX = 1_000_000.0           # por encima: cifra financiera -> descartar
MONTO_REVISAR_MIN = 500.0         # entre MIN y esto: marcar revision
MONTO_REVISAR_MAX = 200_000.0     # entre esto y MAX: marcar revision

# --- PORTAFOLIO: nombres canonicos tal como los quieres en el output ---
PORTAFOLIO = [
    "Cerro Verde", "Pacasmayo", "Aenza", "Ferreycorp", "Engie", "Volcan", "BBVA",
    "Pluz Energia", "Puerto Chancay", "Orygen", "Buenaventura", "Alicorp", "UNACEM",
    "Credicorp", "Nexa Perú", "InRetail", "Auna", "IFS", "Minsur", "Hudbay",
    "Hunt Oil", "Hermes", "Colegios Peruanos", "Casa Andina", "Inca Rail", "Intursa",
    "Jockey Plaza", "Lima Expresa", "Orazen", "Primax", "Rutas de Lima", "Tecsup",
    "Hochshild",
]

# Si True, los archivos que NO mapean a una empresa del portafolio se OMITEN.
# Si False, se procesan igual y salen marcados "[SIN MATCH PORTAFOLIO]".
SOLO_PORTAFOLIO = False
UMBRAL_FUZZY_PORTAFOLIO = 90      # respaldo difuso si no hay alias explicito

# --- Mapa de alias / nombres historicos -> nombre del PORTAFOLIO ---
# Clave en minuscula y sin tildes. Incluye rebrands (Edelnor->Pluz, Edegel->Orygen,
# Enersur->Engie, GyM->Aenza, Milpo/Atacocha->Nexa, Continental->BBVA, etc.).
ALIAS_PORTAFOLIO = {
    "enel distribucion": "Pluz Energia", "enel dx": "Pluz Energia",
    "edelnor": "Pluz Energia", "pluz": "Pluz Energia",
    "enel generacion": "Orygen", "enel gen": "Orygen", "enel gx": "Orygen",
    "edegel": "Orygen", "orygen": "Orygen",
    "enersur": "Engie", "engie": "Engie",
    "grana y montero": "Aenza", "joga gym": "Aenza", "aenza": "Aenza",
    "atacocha": "Nexa Perú", "milpo": "Nexa Perú", "nexa": "Nexa Perú",
    "continental": "BBVA", "bbva": "BBVA",
    "buenaventura": "Buenaventura",
    "cerro verde": "Cerro Verde", "pacasmayo": "Pacasmayo",
    "ferreycorp": "Ferreycorp", "alicorp": "Alicorp", "unacem": "UNACEM",
    "volcan": "Volcan", "credicorp": "Credicorp",
    "inretail": "InRetail", "incarail": "Inca Rail", "inca rail": "Inca Rail",
    "minsur": "Minsur", "auna": "Auna", "hudbay": "Hudbay", "hunt oil": "Hunt Oil",
    "hermes": "Hermes", "colegios": "Colegios Peruanos",
    "casa andina": "Casa Andina", "intursa": "Intursa", "jockey": "Jockey Plaza",
    "lima expresa": "Lima Expresa", "orazen": "Orazen", "primax": "Primax",
    "rutas de lima": "Rutas de Lima", "tecsup": "Tecsup",
    "hochschild": "Hochshild", "puerto chancay": "Puerto Chancay",
    "cosco": "Puerto Chancay", "chancay": "Puerto Chancay",
}
# Claves cortas/ambiguas: solo matchean como TOKEN exacto (no como subcadena).
CLAVES_CORTAS = {
    "bvn": "Buenaventura", "bap": "Credicorp", "ifs": "IFS",
    "gym": "Aenza", "g m": "Aenza", "g&m": "Aenza", "gm": "Aenza",
    "nrp": "Nexa Perú", "hoc": "Hochshild", "hbm": "Hudbay",
    "scco": None, "lds": None,   # explicitamente NO portafolio (None)
}

# Ventana de caracteres alrededor de cada keyword (captura la frase completa
# de la dieta, que suele ser larga).
VENTANA_IZQ = 220
VENTANA_DER = 340

# --- Keywords FUERTES: disparan por si solas ---
KW_FUERTE_STEMS = ["dieta", "retribuc", "remuneraci", "honorari", "emolument"]
# --- Keywords DEBILES: solo si hay contexto de directorio/comite/sesion/dieta ---
KW_DEBIL_STEMS = ["pago", "compensaci", "asignaci", "gratificaci"]
# Regex unico de stems (se evalua sobre texto normalizado alineado)
_RE_KW = re.compile("(" + "|".join(KW_FUERTE_STEMS + KW_DEBIL_STEMS) + ")")
# Contexto requerido para aceptar una keyword debil
_RE_CONTEXTO = re.compile(r"director|comite|sesion|dieta|retribuc|remuneraci")
# "Pista de monto": frases que indican que cerca hay (o debe haber) una cifra
# de dinero. Evita capturar numeros de pagina, articulos o DNI sueltos, y
# evita emitir hallazgos a partir de titulos de agenda ("Fijacion de su
# Retribucion") que no traen importe.
_RE_CUE_MONTO = re.compile(
    r"suma de|asciend|monto de|importe de|dieta de|percib|equivalente a|"
    r"por la suma|pagar como|abonar|fij\w+ como|us\$|s/|dolares|soles"
)
# Filtro ANTI-financiero: si la ventana habla de esto, NO es remuneracion de
# directorio (son utilidades, dividendos, capital, EEFF embebidos, etc.).
_RE_ANTI_FINANCIERO = re.compile(
    r"utilidad|dividendo|resultados acumulados|reserva legal|capital social|"
    r"ventas netas|ingresos por|patrimonio neto|flujo de efectivo|"
    r"impuesto a la renta|aplicacion de resultados|aporte de capital|"
    r"valor nominal|acciones comunes|distribucion de utilidad|prima de capital|"
    r"provision|amortizacion de prestamo|colocacion de bonos"
)

# Palabras de relleno a quitar del nombre de archivo para extraer la empresa
_RE_RELLENO_FILE = re.compile(
    r"\b(jga|joa|jgoa|joaa|ega|egm|eg|acta|junta|obligatoria|anual|general|"
    r"extraordinaria|de|del|la|el|los|accionistas|no|presencial|parte|part|final|"
    r"firmada?|escaneada?|copia|buyback|deslistado|revalorizacion|activos|"
    r"agreement|proposal|directorio|directores|eleccion|asamblea|bonistas|"
    r"obligacionistas|bono|bonos|programa|emision|reapertura|"
    r"enero|febrero|marzo|abril|mayo|junio|julio|agosto|"
    r"setiembre|septiembre|octubre|noviembre|diciembre|v\d+|p\d+|q[1-4]|[1-4]q)\b",
    re.IGNORECASE,
)

# --- Marcadores de pagina de Estados Financieros (para OMITIR) ---
MARCADORES_EEFF = [
    "estado de situacion financiera", "estado de resultados",
    "estado de flujos de efectivo", "estado de cambios en el patrimonio",
    "activo corriente", "pasivo corriente", "patrimonio neto",
    "utilidad bruta", "utilidad operativa", "ingresos por ventas",
    "costo de ventas", "flujo de efectivo y equivalentes",
]

# --- Clasificacion de Concepto (orden importa: 'dieta' primero) ---
PATRONES_CONCEPTO = [
    (r"\bdieta", "Dieta"),
    (r"\bhonorari", "Honorarios"),
    (r"\bemolument", "Emolumentos"),
    (r"\bgratificaci", "Gratificacion"),
    (r"\bcompensaci", "Compensacion"),
    (r"\bretribuc", "Retribucion"),
    (r"\bremuneraci", "Remuneracion"),
    (r"por\s+(?:cada\s+)?(?:sesion|reunion|asistencia)", "Pago por sesion"),
    (r"\bpago\b", "Pago"),
]

# --- Clasificacion de Frecuencia (tolerante a "por cada ...") ---
PATRONES_FRECUENCIA = [
    (r"por\s+(?:cada\s+)?sesion", "Por sesion"),
    (r"por\s+(?:cada\s+)?reunion", "Por sesion"),
    (r"por\s+(?:cada\s+)?asistencia", "Por sesion"),
    (r"\bmensual", "Mensual"),
    (r"al\s+mes\b", "Mensual"),
    (r"\banual", "Anual"),
    (r"al\s+a[nñ]o\b", "Anual"),
    (r"\btrimestral", "Trimestral"),
    (r"unica\s+vez", "Unica vez"),
]

SUFIJOS_SOCIETARIOS = [
    "s.a.a.", "saa", "s.a.c.", "sac", "s.a.", "s a a", "s a c", "s a",
    "s.r.l.", "srl", "e.i.r.l.", "eirl", "compania minera",
    "sociedad anonima abierta", "sociedad anonima cerrada", "sociedad anonima",
]

# Tokens de relleno a quitar del inicio de un nombre de empresa capturado
STOPWORDS_EMPRESA = {
    "acta", "de", "del", "la", "el", "los", "junta", "obligatoria", "anual",
    "accionistas", "no", "presencial", "general", "sociedad", "extraordinaria",
    "y", "fijacion", "su", "designacion", "miembros", "directorio",
}

MESES = {
    "enero": 1, "febrero": 2, "marzo": 3, "abril": 4, "mayo": 5, "junio": 6,
    "julio": 7, "agosto": 8, "setiembre": 9, "septiembre": 9, "octubre": 10,
    "noviembre": 11, "diciembre": 12,
}


# %% ========================= LOGGING / SETUP ===============================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-7s | %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("actas")

if TESSERACT_CMD and os.path.exists(TESSERACT_CMD):
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD
if TESSDATA_DIR:
    # Apunta Tesseract a una carpeta tessdata propia (sin tocar Program Files)
    os.environ["TESSDATA_PREFIX"] = TESSDATA_DIR
os.makedirs(RUTA_CACHE_OCR, exist_ok=True)

# Verificacion temprana de Tesseract y del idioma (unica dependencia externa)
try:
    _ver_tess = pytesseract.get_tesseract_version()
    log.info("Tesseract %s detectado.", _ver_tess)
    try:
        _langs = pytesseract.get_languages(config="")
        if OCR_IDIOMA in _langs:
            log.info("Idioma '%s' disponible.", OCR_IDIOMA)
        else:
            log.error("FALTA el idioma '%s'. Idiomas disponibles: %s. Descarga "
                      "%s.traineddata y ponlo en la carpeta tessdata (o ajusta "
                      "TESSDATA_DIR). Sin esto el OCR no leera nada.",
                      OCR_IDIOMA, _langs, OCR_IDIOMA)
    except Exception as e:
        log.warning("No se pudo listar idiomas de Tesseract: %s", e)
except Exception:
    log.error("Tesseract NO disponible. Instala Tesseract-OCR + idioma 'spa' "
              "y ajusta TESSERACT_CMD. Sin esto el OCR de paginas escaneadas fallara.")


# %% ===================== UTILIDADES DE TEXTO ===============================

def quitar_acentos(texto: str) -> str:
    """Minusculas sin tildes. Puede cambiar la longitud (no usar para indices)."""
    if not texto:
        return ""
    nfkd = unicodedata.normalize("NFKD", texto)
    return "".join(c for c in nfkd if not unicodedata.combining(c)).lower()


def normalizar_alineado(texto: str) -> str:
    """
    Version sin acentos y en minusculas que PRESERVA la longitud (1 char por
    char de entrada). Permite mapear indices del texto normalizado al original
    para recortar la evidencia exacta.
    """
    salida = []
    for ch in texto:
        d = unicodedata.normalize("NFKD", ch)
        base = "".join(c for c in d if not unicodedata.combining(c))
        salida.append((base[:1] or ch).lower())
    return "".join(salida)


def es_pagina_eeff(texto: str) -> bool:
    """True si la pagina parece un Estado Financiero (>=2 marcadores)."""
    n = quitar_acentos(texto)
    hits = sum(1 for m in MARCADORES_EEFF if m in n)
    return hits >= 2


# %% ===================== DETECCION ANIO / TRIMESTRE ========================

_RE_ANIO = re.compile(r"\b(19|20)\d{2}\b")
_RE_TRIM = [
    re.compile(r"\b([1-4])\s*q\s*\d{0,4}\b", re.I),
    re.compile(r"\bq\s*([1-4])\b", re.I),
    re.compile(r"\b([1-4])\s*t\s*\d{0,4}\b", re.I),
    re.compile(r"\bt\s*([1-4])\b", re.I),
]
_TRIM_TEXTO = {"primer": "1", "segundo": "2", "tercer": "3", "cuarto": "4"}
_RE_FECHA_ACTA = re.compile(
    r"(\d{1,2})\s+de\s+([a-zñ]+)\s+(?:de[l]?\s+)?(\d{4})", re.I
)


def detectar_anio_desde_ruta(ruta: str):
    for parte in Path(ruta).parts[:-1]:  # solo carpetas, no el archivo
        m = _RE_ANIO.search(parte)
        if m:
            anio = int(m.group(0))
            if 1995 <= anio <= datetime.now().year + 1:
                return anio
    return None


def detectar_anio_desde_filename(ruta: str):
    m = _RE_ANIO.search(Path(ruta).name)
    if m:
        anio = int(m.group(0))
        if 1995 <= anio <= datetime.now().year + 1:
            return anio
    return None


def detectar_anio_desde_texto(paginas: list) -> int | None:
    """Busca la fecha del acta (ej. '18 de marzo de 2022') en la 1ra pagina."""
    if not paginas:
        return None
    texto = quitar_acentos(paginas[0]["texto"])
    m = _RE_FECHA_ACTA.search(texto)
    if m and m.group(2) in MESES:
        anio = int(m.group(3))
        if 1995 <= anio <= datetime.now().year + 1:
            return anio
    return None


def detectar_trimestre_desde_ruta(ruta: str) -> str:
    for parte in Path(ruta).parts[:-1]:
        p = quitar_acentos(parte)
        for rx in _RE_TRIM:
            m = rx.search(p)
            if m:
                return f"{m.group(1)}Q"
        for palabra, num in _TRIM_TEXTO.items():
            if palabra in p and "trimestre" in p:
                return f"{num}Q"
    return "Sin trimestre"


# %% ============================ OCR ========================================

def _hash_archivo(ruta: str) -> str:
    st = os.stat(ruta)
    clave = f"{ruta}|{st.st_size}|{int(st.st_mtime)}"
    return hashlib.md5(clave.encode("utf-8")).hexdigest()


def _preprocesar_imagen(img_pil: Image.Image) -> Image.Image:
    """Gris -> upscale si es pequeño -> umbral adaptativo -> deskew."""
    if not _HAY_CV2:
        return img_pil.convert("L")
    img = np.array(img_pil.convert("L"))
    h, w = img.shape[:2]
    if w < UMBRAL_UPSCALE_PX:
        img = cv2.resize(img, None, fx=2.0, fy=2.0, interpolation=cv2.INTER_CUBIC)
    img = cv2.bilateralFilter(img, 5, 50, 50)
    img = cv2.adaptiveThreshold(
        img, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 15
    )
    try:
        coords = np.column_stack(np.where(img < 128))
        if len(coords) > 50:
            ang = cv2.minAreaRect(coords)[-1]
            ang = -(90 + ang) if ang < -45 else -ang
            if abs(ang) > 0.5:
                hh, ww = img.shape
                M = cv2.getRotationMatrix2D((ww // 2, hh // 2), ang, 1.0)
                img = cv2.warpAffine(img, M, (ww, hh), flags=cv2.INTER_CUBIC,
                                     borderMode=cv2.BORDER_REPLICATE)
    except Exception:
        pass
    return Image.fromarray(img)


def _texto_y_confianza(data: dict) -> tuple:
    """
    Reconstruye texto CON estructura de lineas a partir de image_to_data,
    agrupando por (bloque, parrafo, linea). Devuelve (texto, confianza_media).
    """
    lineas = {}
    confs = []
    for i in range(len(data["text"])):
        txt = data["text"][i]
        if not txt or not txt.strip():
            continue
        key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
        lineas.setdefault(key, []).append(txt)
        try:
            c = float(data["conf"][i])
            if c >= 0:
                confs.append(c)
        except Exception:
            pass
    texto = "\n".join(" ".join(lineas[k]) for k in sorted(lineas.keys()))
    conf = float(np.mean(confs)) if confs else 0.0
    return texto, round(conf, 1)


def _pixmap_a_pil(page, dpi: int) -> Image.Image:
    """Renderiza una pagina PyMuPDF a imagen PIL al DPI indicado (sin poppler)."""
    zoom = dpi / 72.0
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
    return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)


def ocr_documento(ruta_pdf: str) -> list:
    """
    Devuelve [{'pagina','texto','confianza'}, ...] usando PyMuPDF para render.
    Por cada pagina: si hay capa de texto nativa la usa (confianza=99); si no,
    rasteriza y aplica OCR. Cachea el resultado en disco.
    Lanza excepcion si el PDF esta corrupto (la maneja el orquestador).
    """
    cache_path = os.path.join(RUTA_CACHE_OCR, f"{_hash_archivo(ruta_pdf)}.json")
    if os.path.exists(cache_path):
        try:
            with open(cache_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass

    doc = fitz.open(ruta_pdf)   # FileDataError si esta corrupto
    paginas = []
    fallo_ocr = False
    try:
        for i, page in enumerate(doc, start=1):
            # 1) Intentar capa de texto nativa (PDF no escaneado)
            try:
                nativo = page.get_text("text").strip()
            except Exception:
                nativo = ""
            if len(nativo) >= UMBRAL_TEXTO_NATIVO:
                paginas.append({"pagina": i, "texto": nativo, "confianza": 99.0})
                continue
            # 2) OCR sobre el render
            try:
                img = _pixmap_a_pil(page, OCR_DPI)
                data = pytesseract.image_to_data(
                    _preprocesar_imagen(img), lang=OCR_IDIOMA, config=OCR_CONFIG,
                    output_type=pytesseract.Output.DICT,
                )
                texto, conf = _texto_y_confianza(data)
            except Exception as e:
                log.warning("OCR fallo en pag %d de %s: %s", i, ruta_pdf, e)
                texto, conf = "", 0.0
                fallo_ocr = True
            paginas.append({"pagina": i, "texto": texto, "confianza": conf})
    finally:
        doc.close()

    # Solo se cachea si el OCR fue exitoso y hubo texto. Esto evita envenenar
    # la cache con resultados vacios (p. ej. si Tesseract aun no estaba listo).
    hay_texto = any(p["texto"].strip() for p in paginas)
    if not fallo_ocr and hay_texto:
        try:
            with open(cache_path, "w", encoding="utf-8") as f:
                json.dump(paginas, f, ensure_ascii=False)
        except Exception as e:
            log.warning("No se pudo escribir cache OCR de %s: %s", ruta_pdf, e)
    elif fallo_ocr:
        log.warning("OCR con fallos en %s: no se cachea (se reintentara).", ruta_pdf)
    return paginas


# %% ===================== PARSEO DE MONTOS ==================================

_RE_MONTO = re.compile(
    r"(?P<moneda>US\$|U\$S|S/\.?|\$|soles|d[oó]lares)?\s*"
    r"(?P<num>\d{1,3}(?:[.,]\d{3})+(?:[.,]\d{1,2})?|\d+(?:[.,]\d{1,2})?)",
    re.IGNORECASE,
)


def _parsear_numero(num_str: str):
    s = re.sub(r"[^\d.,]", "", num_str)
    if not s:
        return None
    if "," in s and "." in s:
        if s.rfind(",") > s.rfind("."):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
    elif "," in s:
        partes = s.split(",")
        s = s.replace(",", ".") if len(partes[-1]) == 2 else s.replace(",", "")
    elif "." in s:
        if s.count(".") > 1:
            s = s.replace(".", "")
        elif len(s.split(".")[-1]) == 3:
            s = s.replace(".", "")
    try:
        val = float(s)
        return val if val > 0 else None
    except ValueError:
        return None


def _normalizar_moneda(tok) -> str:
    if not tok:
        return ""
    t = quitar_acentos(tok)
    if t.startswith("us") or t.startswith("u$") or "dolar" in t:
        return "USD"
    if t.startswith("s/") or "sol" in t:
        return "PEN"
    return ""


def extraer_montos(texto: str) -> list:
    """
    Extrae montos plausibles de una ventana de texto.
    - Ignora notacion de centavos '00/100'.
    - Detecta moneda antes (US$/S/) o despues (soles/dolares) del numero.
    - Si hay montos con moneda, descarta los numeros sin moneda (anti-ruido).
    """
    res = []
    for m in _RE_MONTO.finditer(texto):
        start, end = m.span("num")
        despues = texto[end:end + 5]
        antes = texto[max(0, start - 1):start]
        if despues.startswith("/100") or antes == "/":   # guarda de centavos
            continue
        val = _parsear_numero(m.group("num"))
        if val is None:
            continue
        moneda = _normalizar_moneda(m.group("moneda"))
        if not moneda:
            cola = quitar_acentos(texto[end:end + 14])
            if "sol" in cola:
                moneda = "PEN"
            elif "dolar" in cola or "us$" in cola:
                moneda = "USD"
        if not moneda and 1900 <= val <= 2099 and float(val).is_integer():
            continue   # probable anio suelto
        if val < MONTO_MIN or val > MONTO_MAX:
            continue   # numero de pagina/articulo (muy chico) o cifra financiera (enorme)
        res.append({"monto": val, "moneda": moneda, "tiene_moneda": bool(moneda)})

    if any(r["tiene_moneda"] for r in res):
        res = [r for r in res if r["tiene_moneda"]]
    return res


# %% ================ CLASIFICACION TIPO / CONCEPTO / FREC ===================

def clasificar_tipo(ctx_norm: str) -> str:
    if "comite" in ctx_norm:
        return "Comite"
    if "director" in ctx_norm:
        return "Directorio"
    return "Directorio"


def clasificar_concepto(ctx_norm: str) -> str:
    for patron, etiqueta in PATRONES_CONCEPTO:
        if re.search(patron, ctx_norm):
            return etiqueta
    return "Remuneracion"


def clasificar_frecuencia(ctx_norm: str) -> str:
    for patron, etiqueta in PATRONES_FRECUENCIA:
        if re.search(patron, ctx_norm):
            return etiqueta
    return "No especificada"


# %% =================== EXTRACCION POR DOCUMENTO ============================

def extraer_remuneraciones(paginas: list) -> tuple:
    """
    Devuelve (registros, n_paginas_eeff_omitidas).
    Extraccion por VENTANA alrededor de cada keyword, omitiendo paginas EEFF.
    Dedup con preferencia: ante misma (tipo,concepto,monto,pagina) se conserva
    el registro mas informativo (con moneda y frecuencia conocida).
    """
    mejores = {}   # clave -> registro
    eeff_omitidas = 0

    def informatividad(r):
        score = 0
        if r["Moneda"]:
            score += 2
        if r["Frecuencia"] != "No especificada":
            score += 1
        if r["Concepto"] not in ("Remuneracion", "Pago"):
            score += 1
        return score

    for pag in paginas:
        texto = pag["texto"]
        if not texto:
            continue
        if es_pagina_eeff(texto):
            eeff_omitidas += 1
            continue
        conf = pag["confianza"]
        norm = normalizar_alineado(texto)

        for km in _RE_KW.finditer(norm):
            stem = km.group(1)
            i = km.start()
            ini, fin = max(0, i - VENTANA_IZQ), min(len(texto), i + VENTANA_DER)
            ventana = texto[ini:fin]
            ventana_norm = norm[ini:fin]

            es_fuerte = stem in KW_FUERTE_STEMS
            if not es_fuerte and not _RE_CONTEXTO.search(ventana_norm):
                continue   # keyword debil sin contexto -> se descarta
            # Filtro anti-financiero: utilidades/dividendos/capital NO son dietas
            if _RE_ANTI_FINANCIERO.search(ventana_norm):
                continue

            tipo = clasificar_tipo(ventana_norm)
            concepto = clasificar_concepto(ventana_norm)
            frecuencia = clasificar_frecuencia(ventana_norm)
            evidencia = re.sub(r"\s+", " ", ventana).strip()[:300]
            cue = bool(_RE_CUE_MONTO.search(ventana_norm))
            montos = extraer_montos(ventana)
            # Un numero SIN moneda solo se acepta si hay pista de monto cerca
            # (descarta numeros de pagina / articulo / DNI).
            montos = [mh for mh in montos if mh["tiene_moneda"] or cue]

            if not montos:
                # Hallazgo sin cifra: solo si se ESPERABA un importe (pista de
                # monto). Asi se ignoran titulos de agenda como "Fijacion de su
                # Retribucion" que no contienen pago alguno.
                if not cue:
                    continue
                clave = ("HALLAZGO", concepto, None, pag["pagina"])
                reg = {
                    "Tipo": tipo, "Concepto": concepto, "Monto": np.nan,
                    "Moneda": "", "Frecuencia": frecuencia,
                    "Pagina": pag["pagina"], "Confianza_OCR": conf,
                    "Requiere_Revision": True, "Evidencia": evidencia,
                }
                if clave not in mejores:
                    mejores[clave] = reg
                continue

            for mh in montos:
                clave = (tipo, concepto, round(mh["monto"], 2), pag["pagina"])
                fuera_rango = (mh["monto"] < MONTO_REVISAR_MIN
                               or mh["monto"] > MONTO_REVISAR_MAX)
                reg = {
                    "Tipo": tipo, "Concepto": concepto, "Monto": mh["monto"],
                    "Moneda": mh["moneda"], "Frecuencia": frecuencia,
                    "Pagina": pag["pagina"], "Confianza_OCR": conf,
                    "Requiere_Revision": bool(
                        conf < CONFIANZA_FILA_REVISION
                        or not mh["tiene_moneda"] or fuera_rango
                    ),
                    "Evidencia": evidencia,
                }
                if clave not in mejores or informatividad(reg) > informatividad(mejores[clave]):
                    mejores[clave] = reg

    return list(mejores.values()), eeff_omitidas


# %% =================== NOMBRE DE EMPRESA ===================================

# Empresa en mayuscula+minuscula terminada en sufijo societario.
_RE_EMPRESA = re.compile(
    r"([A-ZÁÉÍÓÚÑ][\wÁÉÍÓÚÑáéíóúñ&.\-]*"
    r"(?:\s+[A-ZÁÉÍÓÚÑ][\wÁÉÍÓÚÑáéíóúñ&.\-]*){0,5}"
    r"\s+S\.?\s?A\.?\s?[AC]?\.?)"
)


def empresa_desde_filename(ruta_pdf: str) -> str:
    """
    Limpia el nombre de archivo (patron '[año] JGA/JOA <Empresa>') y devuelve el
    nombre 'crudo' de la empresa. La canonicalizacion al PORTAFOLIO se hace
    despues, en mapear_empresas().
    """
    s = Path(ruta_pdf).stem
    s = re.sub(r"\d{1,2}[.\-/]\d{1,2}[.\-/]\d{2,4}", " ", s)   # fechas dd.mm.aaaa
    s = re.sub(r"\b(19|20)\d{2}(\.\d{1,2})?\b", " ", s)        # años / aaaa.mm
    s = re.sub(r"[A-Za-z]{3}-\d{2}", " ", s)                    # Jul-15
    s = _RE_RELLENO_FILE.sub(" ", s)
    s = re.sub(r"[_\-\.\(\)]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s.title() if len(s) >= 2 else "DESCONOCIDA"


def _a_portafolio(nombre: str):
    """
    Devuelve el nombre del PORTAFOLIO al que corresponde 'nombre', o None.
    Aplica alias/rebrands (Edelnor->Pluz, Edegel->Orygen, GyM->Aenza, etc.).
    """
    n = re.sub(r"[^a-z0-9 ]", " ", quitar_acentos(nombre))
    n = re.sub(r"\s+", " ", n).strip()
    if not n:
        return None
    toks = set(n.split())
    # 1) Alias por subcadena (claves largas primero para evitar choques)
    for clave in sorted(ALIAS_PORTAFOLIO, key=len, reverse=True):
        if clave in n:
            return ALIAS_PORTAFOLIO[clave]
    # 2) Claves cortas/ambiguas: solo como TOKEN exacto (valor puede ser None)
    for clave, canon in CLAVES_CORTAS.items():
        partes = clave.split()
        if (len(partes) == 1 and clave in toks) or \
           (len(partes) > 1 and all(p in toks for p in partes)):
            return canon
    # 3) El nombre ya podria venir como uno del portafolio
    for emp in PORTAFOLIO:
        if quitar_acentos(emp) in n:
            return emp
    return None


def canonizar_empresa(nombre: str) -> str:
    """
    Mapea un nombre crudo al nombre del PORTAFOLIO. Si no hay alias, intenta un
    match difuso contra el portafolio; si tampoco, marca '[SIN MATCH PORTAFOLIO]'
    (no se inventa una equivalencia).
    """
    canon = _a_portafolio(nombre)
    if canon:
        return canon
    base = re.sub(r"[^a-z0-9 ]", " ", quitar_acentos(nombre)).strip()
    mejor, score = None, -1
    for emp in PORTAFOLIO:
        s = fuzz.token_set_ratio(base, quitar_acentos(emp))
        if s > score:
            mejor, score = emp, s
    if score >= UMBRAL_FUZZY_PORTAFOLIO:
        return mejor
    return f"{re.sub(r'\\s+', ' ', nombre).strip().title()} [SIN MATCH PORTAFOLIO]"


def mapear_empresas(nombres_crudos: list) -> dict:
    """Devuelve dict {nombre_crudo -> nombre del PORTAFOLIO (o [SIN MATCH ...])}."""
    mapa = {raw: canonizar_empresa(raw) for raw in set(nombres_crudos)}
    sin = sorted({raw for raw, c in mapa.items() if "[SIN MATCH" in str(c)})
    if sin:
        log.warning("%d nombre(s) fuera del portafolio (revisar/ignorar): %s",
                    len(sin), ", ".join(sin[:20]))
    return mapa


# %% =================== ESCRITURA EXCEL =====================================

COLUMNAS_HOJA = ["Tipo", "Concepto", "Monto", "Moneda", "Frecuencia", "Pagina",
                 "RutaArchivo", "Anio", "Trimestre", "Confianza_OCR",
                 "Requiere_Revision", "Evidencia"]
COLUMNAS_RESUMEN = ["Empresa", "Concepto", "Monto", "Moneda", "Frecuencia",
                    "Pagina", "Anio", "Trimestre"]
COLUMNAS_DIAG = ["RutaArchivo", "Empresa", "Anio", "Trimestre", "Paginas",
                 "Confianza_media", "Paginas_EEFF_omitidas", "Filas_extraidas",
                 "Estado"]


def _sanitizar_hoja(nombre: str, usados: set) -> str:
    n = re.sub(r"[\[\]\:\*\?\/\\]", " ", nombre).strip()[:31] or "Empresa"
    base, i = n, 1
    while n.lower() in usados:
        suf = f" ({i})"; n = base[: 31 - len(suf)] + suf; i += 1
    usados.add(n.lower())
    return n


def _estilo_encabezado(ws, ncols):
    fill = PatternFill("solid", fgColor="1F3864")
    fuente = Font(bold=True, color="FFFFFF", size=10)
    borde = Border(bottom=Side(style="thin", color="999999"))
    for c in range(1, ncols + 1):
        cel = ws.cell(row=1, column=c)
        cel.fill = fill; cel.font = fuente; cel.border = borde
        cel.alignment = Alignment(vertical="center", horizontal="center")
    ws.freeze_panes = "A2"


def _autoancho(ws, df, max_ancho=60):
    for j, col in enumerate(df.columns, start=1):
        largo = max([len(str(col))] +
                    [len(str(v)) for v in df.iloc[:, j - 1].head(200)])
        ws.column_dimensions[get_column_letter(j)].width = min(largo + 2, max_ancho)


def _volcar_df(ws, df, columnas):
    ws.append(columnas)
    for _, fila in df[columnas].iterrows():
        ws.append([("" if pd.isna(v) else v) for v in fila.tolist()])
    _estilo_encabezado(ws, len(columnas))
    _autoancho(ws, df[columnas])


def escribir_excel(df: pd.DataFrame, df_diag: pd.DataFrame, ruta_salida: str):
    wb = Workbook(); wb.remove(wb.active)

    # RESUMEN (anio mas reciente con datos validos por Empresa+Concepto)
    val = df[df["Monto"].notna() & (~df["Requiere_Revision"])].copy()
    if val.empty:
        val = df[df["Monto"].notna()].copy()
    if not val.empty:
        val["Anio"] = pd.to_numeric(val["Anio"], errors="coerce")
        val = val.sort_values(["Empresa", "Concepto", "Anio"],
                              ascending=[True, True, False])
        resumen = val.drop_duplicates(subset=["Empresa", "Concepto"],
                                      keep="first")[COLUMNAS_RESUMEN]
    else:
        resumen = pd.DataFrame(columns=COLUMNAS_RESUMEN)
    _volcar_df(wb.create_sheet("RESUMEN"), resumen, COLUMNAS_RESUMEN)

    # DIAGNOSTICO (todos los archivos, hayan dado datos o no)
    _volcar_df(wb.create_sheet("DIAGNOSTICO"), df_diag, COLUMNAS_DIAG)

    # Una hoja por empresa
    usados = {"resumen", "diagnostico"}
    rojo = PatternFill("solid", fgColor="FCE4E4")
    for empresa, sub in df.groupby("Empresa", sort=True):
        ws = wb.create_sheet(_sanitizar_hoja(str(empresa), usados))
        _volcar_df(ws, sub, COLUMNAS_HOJA)
        col_rev = COLUMNAS_HOJA.index("Requiere_Revision") + 1
        for r in range(2, ws.max_row + 1):
            if ws.cell(row=r, column=col_rev).value in (True, "True", "VERDADERO"):
                for c in range(1, len(COLUMNAS_HOJA) + 1):
                    ws.cell(row=r, column=c).fill = rojo

    wb.save(ruta_salida)
    log.info("Excel generado: %s", ruta_salida)


# %% =================== ORQUESTADOR PRINCIPAL ===============================

def descubrir_archivos(raiz: str) -> list:
    """
    PDFs + Word (.docx/.doc) bajo la raiz, ignorando la cache.
    Solo incluye archivos cuya carpeta de año este en ANIOS_INCLUIR
    (editable arriba). Archivos sin año de carpeta reconocible se omiten.
    """
    exts = (".pdf", ".docx", ".doc")
    rutas = []
    for dirpath, _dirs, files in os.walk(raiz):
        if "_ocr_cache" in dirpath:
            continue
        for f in files:
            if f.startswith("~$"):           # temporales de Office
                continue
            if not f.lower().endswith(exts):
                continue
            ruta = os.path.join(dirpath, f)
            anio = detectar_anio_desde_ruta(ruta)
            if anio is None or anio not in ANIOS_INCLUIR:
                continue                      # fuera del rango de años pedido
            rutas.append(ruta)
    return sorted(rutas)


def _contar_paginas_pdf(ruta: str) -> int:
    """Conteo de paginas instantaneo con PyMuPDF (sin OCR)."""
    d = fitz.open(ruta)
    n = d.page_count
    d.close()
    return n


def leer_docx(ruta: str, parrafos_por_pagina: int = 25) -> list:
    """
    Lee un .docx (parrafos + tablas) y lo devuelve como 'paginas' del mismo
    formato que el OCR, troceado para que el filtro anti-EEFF actue localmente.
    """
    if not _HAY_DOCX:
        raise RuntimeError("python-docx no instalado (pip install python-docx)")
    d = docx.Document(ruta)
    bloques = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            celdas = [c.text.strip() for c in row.cells if c.text.strip()]
            if celdas:
                bloques.append("  ".join(celdas))
    paginas = []
    for i in range(0, len(bloques), parrafos_por_pagina):
        chunk = "\n".join(bloques[i:i + parrafos_por_pagina])
        paginas.append({"pagina": i // parrafos_por_pagina + 1,
                        "texto": chunk, "confianza": 100.0})
    return paginas or [{"pagina": 1, "texto": "", "confianza": 0.0}]


def procesar_un_archivo(ruta: str) -> tuple:
    """
    Worker (apto para hilos): procesa UN archivo y devuelve (filas, fila_diag).
    Aplica el filtro de paginas para PDFs y maneja Word. Nunca lanza: los
    errores se reportan en el estado del diagnostico.
    """
    ext = Path(ruta).suffix.lower()
    nombre = os.path.basename(ruta)
    trimestre = detectar_trimestre_desde_ruta(ruta)
    empresa_raw = empresa_desde_filename(ruta)
    diag = {"RutaArchivo": ruta, "Empresa": empresa_raw, "Anio": "",
            "Trimestre": trimestre, "Paginas": 0, "Confianza_media": 0.0,
            "Paginas_EEFF_omitidas": 0, "Filas_extraidas": 0, "Estado": "OK"}

    # Filtro opcional: si solo interesa el portafolio, descartar lo que no mapea
    if SOLO_PORTAFOLIO and "[SIN MATCH" in canonizar_empresa(empresa_raw):
        diag["Estado"] = "OMITIDO_FUERA_PORTAFOLIO"
        return [], diag

    try:
        if ext == ".pdf":
            n_pag = _contar_paginas_pdf(ruta)
            diag["Paginas"] = n_pag
            # Filtro de paginas opcional (MIN_PAGINAS_ACTA=0 => desactivado).
            # Salta el descarte si el nombre dice "parte" (acta en varios archivos).
            if (MIN_PAGINAS_ACTA and n_pag < MIN_PAGINAS_ACTA
                    and not re.search(r"\bpart", nombre, re.I)):
                diag["Estado"] = "OMITIDO_CORTO"
                return [], diag
            paginas = ocr_documento(ruta)
        elif ext == ".docx":
            if not PROCESAR_WORD:
                diag["Estado"] = "OMITIDO_WORD"; return [], diag
            paginas = leer_docx(ruta)
            diag["Paginas"] = len(paginas)
        elif ext == ".doc":
            # .doc binario antiguo: python-docx no lo lee
            diag["Estado"] = "OMITIDO_DOC_BINARIO"
            return [], diag
        else:
            diag["Estado"] = "OMITIDO"; return [], diag
    except Exception as e:
        log.error("Lectura/OCR fallida: %s | %s", ruta, e)
        diag["Estado"] = "ERROR_OCR"
        return [], diag

    diag["Anio"] = (detectar_anio_desde_ruta(ruta)
                    or detectar_anio_desde_filename(ruta)
                    or detectar_anio_desde_texto(paginas) or "")

    if not (paginas and any(p["texto"] for p in paginas)):
        diag["Estado"] = "OCR_VACIO"
        return [], diag

    diag["Confianza_media"] = round(
        float(np.mean([p["confianza"] for p in paginas])), 1)
    try:
        registros, eeff = extraer_remuneraciones(paginas)
    except Exception as e:
        log.error("Error de extraccion en %s | %s", ruta, e)
        diag["Estado"] = "ERROR_EXTRACCION"
        return [], diag

    diag["Paginas_EEFF_omitidas"] = eeff
    diag["Filas_extraidas"] = len(registros)
    diag["Estado"] = "OK" if registros else "SIN_REMUNERACION"
    filas = []
    for r in registros:
        r.update({"Empresa_raw": empresa_raw, "RutaArchivo": ruta,
                  "Anio": diag["Anio"], "Trimestre": trimestre})
        filas.append(r)
    return filas, diag


def procesar_todo(raiz: str = RUTA_RAIZ, ruta_excel: str = RUTA_SALIDA_EXCEL):
    archivos = descubrir_archivos(raiz)
    n_pdf = sum(1 for a in archivos if a.lower().endswith(".pdf"))
    n_doc = len(archivos) - n_pdf
    log.info("Encontrados %d archivos (%d PDF, %d Word) bajo %s",
             len(archivos), n_pdf, n_doc, raiz)
    log.info("Procesando en paralelo con %d hilos...", MAX_WORKERS)

    filas, diag, errores = [], [], []
    total = len(archivos)
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as ex:
        futs = {ex.submit(procesar_un_archivo, r): r for r in archivos}
        for k, fut in enumerate(as_completed(futs), start=1):
            ruta = futs[fut]
            try:
                f_rows, d_row = fut.result()
            except Exception as e:   # red de seguridad: el worker no deberia lanzar
                f_rows, d_row = [], {
                    "RutaArchivo": ruta, "Empresa": empresa_desde_filename(ruta),
                    "Anio": "", "Trimestre": detectar_trimestre_desde_ruta(ruta),
                    "Paginas": 0, "Confianza_media": 0.0,
                    "Paginas_EEFF_omitidas": 0, "Filas_extraidas": 0,
                    "Estado": "ERROR_OCR"}
            filas.extend(f_rows)
            diag.append(d_row)
            if d_row["Estado"].startswith("ERROR"):
                errores.append({"ruta": ruta, "error": d_row["Estado"]})
            log.info("[%d/%d] %s | %s | filas=%d", k, total,
                     os.path.basename(ruta), d_row["Estado"], d_row["Filas_extraidas"])

    df_diag = pd.DataFrame(diag, columns=COLUMNAS_DIAG)

    # Canonicalizar nombres de empresa (maestro-first) sobre TODOS los archivos,
    # para que las hojas, el RESUMEN y el DIAGNOSTICO usen el mismo nombre.
    mapa = mapear_empresas(df_diag["Empresa"].tolist()) if not df_diag.empty else {}
    df_diag["Empresa"] = df_diag["Empresa"].map(lambda r: mapa.get(r, str(r).title()))

    if not filas:
        log.warning("No se extrajo ningun dato. Se genera Excel solo con DIAGNOSTICO.")
        df = pd.DataFrame(columns=["Empresa"] + COLUMNAS_HOJA)
    else:
        df = pd.DataFrame(filas)
        df["Empresa"] = df["Empresa_raw"].map(lambda r: mapa.get(r, str(r).title()))

    escribir_excel(df, df_diag, ruta_excel)

    if errores:
        ruta_err = os.path.splitext(ruta_excel)[0] + "_ERRORES.csv"
        pd.DataFrame(errores).to_csv(ruta_err, index=False, encoding="utf-8-sig")
        log.info("Errores -> %s (%d archivos)", ruta_err, len(errores))

    log.info("=" * 60)
    log.info("Filas extraidas:        %d", len(df))
    if not df.empty:
        log.info("Empresas consolidadas:  %d", df["Empresa"].nunique())
        log.info("Filas a revisar:        %d", int(df["Requiere_Revision"].sum()))
    vc = df_diag["Estado"].value_counts()
    for estado, cnt in vc.items():
        log.info("  %-20s %d", estado, cnt)
    log.info("=" * 60)
    return df, df_diag


# %% ============================ EJECUCION ==================================

if __name__ == "__main__":
    df_resultado, df_diagnostico = procesar_todo()